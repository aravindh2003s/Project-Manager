from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

# ── Helper: add horizontal rule ───────────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2C3E6B')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ── Helper: section heading ───────────────────────────────────────────────────
def add_section_heading(doc, text):
    add_hr(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)
    return p

# ── Helper: bullet point ──────────────────────────────────────────────────────
def add_bullet(doc, text, indent=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    return p

# ── Helper: normal paragraph ──────────────────────────────────────────────────
def add_normal(doc, text, bold=False, size=9.5, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# NAME & CONTACT
# ═══════════════════════════════════════════════════════════════════════════════
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_before = Pt(0)
name_p.paragraph_format.space_after  = Pt(2)
name_run = name_p.add_run("ARAVINDHAN S")
name_run.bold           = True
name_run.font.size      = Pt(20)
name_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after  = Pt(3)
title_run = title_p.add_run("QA Automation Engineer | Selenium | Playwright | TOSCA (Learning)")
title_run.bold           = True
title_run.font.size      = Pt(10.5)
title_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_before = Pt(0)
contact_p.paragraph_format.space_after  = Pt(4)
contact_run = contact_p.add_run(
    "aravindh2003s@gmail.com  |  8838544167  |  "
    "linkedin.com/in/aravindhan-s-37254b2a2  |  github.com/aravindh2003s"
)
contact_run.font.size      = Pt(9)
contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ═══════════════════════════════════════════════════════════════════════════════
# PROFESSIONAL SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Professional Summary")
summary = (
    "QA Automation Engineer with hands-on experience in building and executing automated test frameworks "
    "using Selenium, Playwright, TypeScript, and Java. "
    "Experienced in test planning, test design, execution, and reporting in Agile environments. "
    "Currently upskilling in Tricentis TOSCA and model-based test automation. "
    "Strong in API testing, CI/CD integration, and defect analysis. "
    "Passionate about delivering high-quality software and improving test efficiency."
)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run(summary)
run.font.size = Pt(9.5)

# ═══════════════════════════════════════════════════════════════════════════════
# WORK EXPERIENCE
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Work Experience")

# Job title row
job_p = doc.add_paragraph()
job_p.paragraph_format.space_before = Pt(4)
job_p.paragraph_format.space_after  = Pt(1)
job_run = job_p.add_run("Automation Test Engineer Intern")
job_run.bold      = True
job_run.font.size = Pt(10)
job_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)

# Company + dates
comp_p = doc.add_paragraph()
comp_p.paragraph_format.space_before = Pt(0)
comp_p.paragraph_format.space_after  = Pt(2)
comp_run = comp_p.add_run("Larsen & Toubro  |  Chennai  |  02/2025 – 05/2025")
comp_run.font.size      = Pt(9.5)
comp_run.italic         = True
comp_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

bullets = [
    "Designed and executed automated test scripts for web applications using Playwright and Selenium",
    "Participated in test planning, test design, and test execution across multiple sprint cycles",
    "Generated and published automation test reports for stakeholder review",
    "Performed failure analysis and defect reporting using structured incident management practices",
    "Integrated automated tests into CI/CD pipeline, reducing manual testing effort by 40%",
    "Worked in Agile/Scrum environment; participated in sprint planning, standups, and retrospectives",
    "Gained exposure to model-based testing concepts aligned with Tricentis TOSCA methodology",
    "Built reusable automation frameworks with Cucumber BDD, enhancing test maintainability",
]
for b in bullets:
    add_bullet(doc, b)

# ═══════════════════════════════════════════════════════════════════════════════
# SKILLS
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Technical Skills")

skill_groups = [
    ("Automation Testing",    "Selenium WebDriver, Playwright, Cucumber BDD, Tricentis TOSCA (Learning)"),
    ("Testing Activities",    "Test Planning, Test Design, Test Execution, Test Reporting, Failure Analysis, Defect Tracking, Incident Management"),
    ("API Testing",           "REST API Testing, Postman"),
    ("Programming",           "Java (Core Java), TypeScript, JavaScript"),
    ("Frontend",              "HTML, CSS"),
    ("Backend / Database",    "FastAPI, REST APIs, SQL (MySQL)"),
    ("Tools & Practices",     "CI/CD Pipelines, Git/GitHub, Agile/Scrum, SDLC"),
]

for label, value in skill_groups:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    bold_run = p.add_run(f"{label}: ")
    bold_run.bold      = True
    bold_run.font.size = Pt(9.5)
    bold_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)
    val_run = p.add_run(value)
    val_run.font.size = Pt(9.5)

# ═══════════════════════════════════════════════════════════════════════════════
# CERTIFICATIONS
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Certifications")
add_bullet(doc, "Tricentis TOSCA Automation – In Progress (2026)")

# ═══════════════════════════════════════════════════════════════════════════════
# PROJECTS
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Projects")

projects = [
    (
        "Web Application Automation Framework",
        [
            "Designed automated test scripts for UI workflows using Playwright and Selenium WebDriver",
            "Executed regression and smoke test suites across multiple environments",
            "Generated automation test execution reports documenting pass/fail rates and defect trends",
        ]
    ),
    (
        "E-Commerce Application Automation Testing",
        [
            "Automated checkout, payment, and cart workflows for a full-stack e-commerce platform",
            "Performed defect reporting and failure analysis on identified test failures",
            "Executed cross-browser test automation to ensure compatibility across browsers",
        ]
    ),
    (
        "Employee Management System – Automation Testing",
        [
            "Created test scenarios and automated test cases for CRUD operations on employee records",
            "Performed database validation using SQL and API testing using Postman",
            "Integrated test suite into CI/CD pipeline for continuous quality assurance",
        ]
    ),
    (
        "Indoor Positioning System – AI & IoT Based",
        [
            "AI-based solution to detect and track indoor objects",
            "Implemented ML algorithms for enhanced tracking and navigation accuracy",
        ]
    ),
]

for proj_title, proj_bullets in projects:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(proj_title)
    run.bold      = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)
    for b in proj_bullets:
        add_bullet(doc, b)

# ═══════════════════════════════════════════════════════════════════════════════
# EDUCATION
# ═══════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, "Education")

edu_entries = [
    ("Sri Venkateshwaraa College of Engineering and Technology",
     "Bachelor of Technology – Computer Science and Engineering",
     "2021 – 2025  |  Puducherry, India"),
    ("Jawahar Higher Secondary School",
     "Class XII (Senior Secondary)",
     "2020 – 2021  |  Puducherry, India"),
    ("Jawahar Higher Secondary School",
     "Class X (Secondary School)",
     "2018 – 2019  |  Puducherry, India"),
]

for inst, deg, period in edu_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(inst)
    run.bold      = True
    run.font.size = Pt(9.5)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run(deg)
    r2.font.size = Pt(9.5)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(2)
    r3 = p3.add_run(period)
    r3.font.size      = Pt(9)
    r3.italic         = True
    r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path     = r'c:\Users\lovel\OneDrive\Desktop\Resume\Assets\ARAVINDHAN_Resume_TOSCA_Updated.docx'
pdf_out_path = r'c:\Users\lovel\OneDrive\Desktop\Resume\Assets\ARAVINDHAN_Resume_TOSCA_Updated.pdf'

doc.save(out_path)
print(f"Saved DOCX: {out_path}")

try:
    print("Converting to PDF...")
    convert(out_path, pdf_out_path)
    print(f"Saved PDF:  {pdf_out_path}")
except Exception as e:
    print(f"PDF conversion failed: {e}")
    print("Note: docx2pdf requires Microsoft Word to be installed for PDF conversion.")

