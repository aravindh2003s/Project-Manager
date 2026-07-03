"""
Build ARAVINDHAN S – Software Developer Resume (DOCX + PDF)

Formatting rules:
  • Font       : Calibri everywhere, Name in Calibri Bold 18pt
  • Body text  : 10.5 pt
  • Headings   : 12 pt Bold ALL CAPS + bottom border
  • Margins    : Narrow (1.27 cm / 0.5 in all sides)
  • Spacing    : Line spacing 1.15, Before 0 pt, After 6 pt
  • Bullets    : • (solid dot)
  • PDF export : Word ExportAsFixedFormat (not Print-to-PDF)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Constants ─────────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x3A, 0x6B)
BLUE = RGBColor(0x1A, 0x73, 0xE8)
BLK  = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
FONT = 'Calibri'

doc = Document()

# ── Page Setup: Narrow margins (1.27 cm = 0.5 in) ────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin   = Cm(1.8)
    sec.right_margin  = Cm(1.8)

# ── Set default document style ────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(10.5)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after  = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Fix font for East Asian / complex scripts (ensures Calibri everywhere)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:ascii'),    FONT)
rFonts.set(qn('w:hAnsi'),    FONT)
rFonts.set(qn('w:eastAsia'), FONT)
rFonts.set(qn('w:cs'),       FONT)

# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def set_run(run, size=10.5, bold=False, italic=False, color=BLK):
    """Apply consistent font settings to a run."""
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure Calibri is set in XML too (avoids substitution in PDF)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'),    FONT)
    rFonts.set(qn('w:hAnsi'),    FONT)
    rFonts.set(qn('w:eastAsia'), FONT)
    rFonts.set(qn('w:cs'),       FONT)

def set_spacing(p, before=0, after=6, line=1.15):
    """Set paragraph spacing."""
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = line

def add_section_heading(text):
    """12pt Bold ALL CAPS with bottom border — professional section divider."""
    p = doc.add_paragraph()
    set_spacing(p, before=18, after=10, line=1.15)
    run = p.add_run(text.upper())
    set_run(run, size=12, bold=True, color=NAVY)

    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')       # thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1B3A6B')  # navy
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_bullet(text):
    """Body bullet using • (solid dot) with hanging indent."""
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=2, line=1.15)
    p.paragraph_format.left_indent   = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    run = p.add_run(f'•  {text}')
    set_run(run, size=11, color=BLK)
    return p

def add_skill_row(label, value):
    """Bold label: value on a single line."""
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=6, line=1.15)
    lbl = p.add_run(f'{label}: ')
    set_run(lbl, size=10.5, bold=True, color=NAVY)
    val = p.add_run(value)
    set_run(val, size=11, color=BLK)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# RESUME CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

# ── NAME ──────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, before=0, after=2, line=1.15)
run = p.add_run('ARAVINDHAN S')
set_run(run, size=18, bold=True, color=NAVY)

# ── TITLE ─────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, before=0, after=2, line=1.15)
run = p.add_run('Frontend Developer (React | TypeScript | JavaScript)')
set_run(run, size=11, bold=True, color=BLUE)

# ── CONTACT (single line with | separators) ──────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, before=0, after=6, line=1.15)
run = p.add_run(
    'aravindh2003s@gmail.com  |  +91 88385 44167  |  '
    'linkedin.com/in/aravindhan-s-37254b2a2  |  github.com/aravindh2003s'
)
set_run(run, size=11, color=GREY)


# ── PROFESSIONAL SUMMARY ─────────────────────────────────────────────────────
add_section_heading('Professional Summary')

p = doc.add_paragraph()
set_spacing(p, before=15, after=15, line=1.2)
run = p.add_run(
    'Frontend Developer with hands-on experience building responsive, scalable web applications '
    'using React, TypeScript, JavaScript, HTML and CSS. Strong understanding of UI architecture, '
    'REST API integration and component-based development. Experienced working in Agile teams, '
    'CI/CD environments and UI automation testing. Passionate about creating clean, user-friendly '
    'interfaces and writing maintainable code.'
)
set_run(run, size=11, color=BLK)


# ── TECH STACK ────────────────────────────────────────────────────────────────
add_section_heading('Tech Stack')

add_skill_row('Frontend',  'React.js, JavaScript (ES6+), TypeScript, HTML5, CSS3, Responsive Design')
add_skill_row('Backend',   'FastAPI, Node.js, Express.js, REST APIs')
add_skill_row('Database',  'MySQL, MongoDB')
add_skill_row('Testing',   'Playwright, Selenium, Cucumber')
add_skill_row('Tools',     'Git, GitHub, CI/CD, Agile/Scrum')


# ── CORE FRONTEND SKILLS (compact inline) ─────────────────────────────────────
add_section_heading('Core Frontend Skills')

p = doc.add_paragraph()
set_spacing(p, before=15, after=15, line=1.2)
run = p.add_run(
    'Component-Based Architecture & Reusable UI  •  '
    'State Management (Context API)  •  '
    'Cross-Browser Compatibility  •  '
    'API Integration & Async Data Handling  •  '
    'Performance Optimization  •  '
    'Mobile-First Responsive Design'
)
set_run(run, size=11, color=BLK)


# ── EXPERIENCE ────────────────────────────────────────────────────────────────
add_section_heading('Experience')

# Job title + company
p = doc.add_paragraph()
set_spacing(p, before=0, after=2, line=1.15)
r1 = p.add_run('Automation Tester Intern — Larsen & Toubro, Chennai')
set_run(r1, size=11, bold=True, color=NAVY)
r2 = p.add_run('     Feb 2025 – May 2025')
set_run(r2, size=10.5, italic=True, color=GREY)

for b in [
    'Developed automated UI test scripts using Playwright + TypeScript for critical user workflows',
    'Built reusable automation framework using Cucumber BDD, improving test maintainability',
    'Collaborated with developers to identify UI issues early and improve product quality',
    'Integrated automated testing into CI/CD pipelines, reducing manual testing effort by 40%',
    'Performed cross-browser testing to ensure consistent UI behavior',
]:
    add_bullet(b)


# ── PROJECTS ──────────────────────────────────────────────────────────────────
add_section_heading('Projects')

projects = [
    ('E-Commerce Custom T-Shirt Platform (React)', [
        'Developed product browsing, cart and checkout experience',
        'Built dynamic product customization interface',
        'Focused on UX, responsiveness and performance optimization',
    ]),
    ('Employee Management System (React + FastAPI + MySQL)', [
        'Built CRUD dashboard UI to manage employee records',
        'Implemented form validation and API error handling',
    ]),
    ('Portfolio Website (React)', [
        'Built responsive portfolio with authentication and dynamic project display',
    ]),
]

for proj_title, proj_bullets in projects:
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=2, line=1.15)
    run = p.add_run(proj_title)
    set_run(run, size=11, bold=True, color=NAVY)
    for b in proj_bullets:
        add_bullet(b)


# ── EDUCATION ─────────────────────────────────────────────────────────────────
add_section_heading('Education')

edu = [
    ('Sri Venkateshwaraa College of Engineering and Technology',
     'B.Tech – Computer Science and Engineering', '2021 – 2025'),
    ('Jawahar Higher Secondary School',
     'Class XII – Higher Secondary', '2020 – 2021'),
    ('Jawahar Higher Secondary School',
     'Class XI – Secondary School', '2018 – 2019'),
]

for inst, deg, period in edu:
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=0, line=1.15)
    run = p.add_run(inst)
    set_run(run, size=10.5, bold=True, color=BLK)

    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=0, line=1.15)
    r2 = p2.add_run(f'{deg}  •  {period}')
    set_run(r2, size=10.5, italic=True, color=GREY)


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
out_path = r'c:\Users\lovel\OneDrive\Desktop\Resume\Assets\ARAVINDHAN_SoftwareDeveloper_Resume.docx'
pdf_out  = r'c:\Users\lovel\OneDrive\Desktop\Resume\Assets\ARAVINDHAN_SoftwareDeveloper_Resume.pdf'

doc.save(out_path)
print(f'Saved DOCX: {out_path}')

# ── Export to PDF via Word COM (File → Save As → PDF) ─────────────────────────
# This uses ExportAsFixedFormat which is the "Save As PDF" method,
# NOT the Print-to-PDF path that breaks spacing.
try:
    import win32com.client
    abs_docx = os.path.abspath(out_path)
    abs_pdf  = os.path.abspath(pdf_out)

    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False

    doc_com = word.Documents.Open(abs_docx)
    doc_com.ExportAsFixedFormat(
        OutputFileName=abs_pdf,
        ExportFormat=17,        # wdExportFormatPDF
        OpenAfterExport=False,
        OptimizeFor=0,          # wdExportOptimizeForPrint (best quality)
        Range=0,                # wdExportAllDocument
        Item=0,                 # wdExportDocumentContent
        IncludeDocProps=True,
    )
    doc_com.Close(False)
    word.Quit()
    print(f'Saved PDF:  {abs_pdf}')
    print('(Exported via Word -> Save As -> PDF -- best quality)')

except ImportError:
    print('pywin32 not installed -- skipping PDF export.')
    print('Open the DOCX in Word -> File -> Save As -> PDF (Best for electronic distribution)')
except Exception as e:
    print(f'Word COM unavailable -- skipping PDF export: {e}')
    print('Open the DOCX in Word -> File -> Save As -> PDF (Best for electronic distribution)')
